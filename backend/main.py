import os
import uuid
import shutil
import logging
import json
import asyncio
import time
from typing import Optional
from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()

from pydantic import BaseModel
from fastapi import FastAPI, UploadFile, File, HTTPException, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, StreamingResponse, HTMLResponse
from fastapi.staticfiles import StaticFiles

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s"
)
logger = logging.getLogger("backend.main")

import re

# Import handlers using robust import structure
try:
    from backend.docx_processor import DocxProcessor
    from backend.metrics import get_text_metrics
    from backend.humanizer import FrenchHumanizer, RateLimitError, GeminiAPIError, HumanizerError
    from backend.post_processor import FrenchPostProcessor
except ImportError:
    from docx_processor import DocxProcessor
    from metrics import get_text_metrics
    from humanizer import FrenchHumanizer, RateLimitError, GeminiAPIError, HumanizerError
    from post_processor import FrenchPostProcessor

def calculate_metrics(text: str) -> dict:
    if not text:
        text = ""
    res = get_text_metrics(text)
    ai_prob = res.get("ai_probability_score", 0.0)
    return {
        "word_count": res.get("word_count", 0),
        "character_count": res.get("character_count", 0),
        "sentence_count": res.get("sentence_count", 0),
        "syllable_count": res.get("syllable_count", 0),
        "readability_score": res.get("readability_index", 0.0),
        "ai_score": ai_prob,
        "human_score": round(100.0 - ai_prob, 1),
        "grade_level": res.get("readability_grade", "N/A"),
        "vocabulary_richness": res.get("vocabulary_richness_ttr", 0.0),
        "estimated_reading_time_minutes": res.get("estimated_reading_time_minutes", 0.0),
        "estimated_reading_time_seconds": res.get("estimated_reading_time_seconds", 0)
    }

def calculate_paragraph_risks(text: str) -> list:
    if not text:
        return []
    paragraphs = [p.strip() for p in text.split("\n\n")]
    paragraphs = [p for p in paragraphs if p]
    
    results = []
    for p in paragraphs:
        res = get_text_metrics(p)
        ai_score = res.get("ai_probability_score", 0.0)
        plag_score = res.get("plagiarism_risk_score", 0.0)
        max_score = max(ai_score, plag_score)
        
        if max_score > 60:
            risk_level = "high"
        elif max_score > 30:
            risk_level = "medium"
        else:
            risk_level = "low"
            
        results.append({
            "ai_score": ai_score,
            "plagiarism_score": plag_score,
            "risk_level": risk_level
        })
    return results

async def humanize_text(text: str, api_key: Optional[str] = None, preset: str = "balanced", ngram_shield: bool = True) -> str:
    post_processor = FrenchPostProcessor(ngram_shield=ngram_shield)
    key = api_key or os.environ.get("GEMINI_API_KEY")
    if not key:
        logger.warning("No Gemini API key provided. Using local rule-based humanizer.")
        logger.info("Applying FrenchPostProcessor for local post-processing...")
        return post_processor.process(text, ngram_shield=ngram_shield)

    humanizer = FrenchHumanizer(api_key=key, batch_size=2, preset=preset)
    result = await humanizer.humanize_text(text)
    logger.info("Applying anti-detection post-processing...")
    return post_processor.process(result, ngram_shield=ngram_shield)

# Initialize FastAPI application
app = FastAPI(
    title="SAVEYOURDOCUMENT - Document Humanization API",
    description="SAVEYOURDOCUMENT API by Mitrixo Systems - Backend service for AI document humanization and plagiarism removal."
)

# 1. CORS Middleware configuration
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:5173",
        "http://127.0.0.1:5173",
        "http://localhost:3000",
        "*"  # Allow Cloud Run domains
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Temporary upload folder configuration inside backend/temp_uploads
TEMP_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "temp_uploads")
os.makedirs(TEMP_DIR, exist_ok=True)
logger.info(f"Temporary uploads directory initialized at {TEMP_DIR}")

async def cleanup_temp_files_loop():
    """
    Background task that periodically cleans up old files in TEMP_DIR.
    Runs every hour, deleting files older than 1 hour.
    """
    logger.info("Temporary files auto-cleanup background service started.")
    while True:
        try:
            now = time.time()
            # 1 hour = 3600 seconds
            retention_period = 3600
            
            logger.info("Running scheduled cleanup scan in temp_uploads...")
            count = 0
            for filename in os.listdir(TEMP_DIR):
                # Do not delete the .gitkeep file
                if filename == ".gitkeep":
                    continue
                    
                file_path = os.path.join(TEMP_DIR, filename)
                if os.path.isfile(file_path):
                    file_mtime = os.path.getmtime(file_path)
                    if now - file_mtime > retention_period:
                        try:
                            os.remove(file_path)
                            count += 1
                            logger.info(f"Cleaned up expired temp file: {filename}")
                        except Exception as err:
                            logger.error(f"Failed to delete {filename}: {err}")
            if count > 0:
                logger.info(f"Cleanup scan finished. Removed {count} expired files.")
            else:
                logger.info("Cleanup scan finished. No expired files found.")
                
        except Exception as e:
            logger.error(f"Error in cleanup_temp_files_loop: {e}")
            
        # Wait for 1 hour (3600 seconds) before scanning again
        await asyncio.sleep(3600)

@app.on_event("startup")
async def startup_event():
    asyncio.create_task(cleanup_temp_files_loop())

# Schema for humanization request
class HumanizeRequest(BaseModel):
    file_id: str
    api_key: Optional[str] = None
    preset: Optional[str] = "balanced"
    ngram_shield: Optional[bool] = True

@app.get("/")
def read_root():
    static_index = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "static", "index.html")
    if os.path.isfile(static_index):
        return FileResponse(static_index, media_type="text/html")
    return {"status": "healthy", "service": "SAVEYOURDOCUMENT API by Mitrixo Systems"}

# 2. POST /api/upload
@app.post("/api/upload")
async def upload_file(file: UploadFile = File(...)):
    logger.info("Step 1: Document upload request received.")
    logger.info(f"Filename: '{file.filename}', Content Type: '{file.content_type}'")
    
    # Validate extension
    if not file.filename.endswith(".docx"):
        logger.warning(f"Validation failure: Rejected file with invalid extension '{file.filename}'.")
        raise HTTPException(
            status_code=400, 
            detail="Only word documents (.docx) are supported."
        )
    logger.info("Step 2: Extension validation passed (.docx format confirmed).")
        
    file_id = str(uuid.uuid4())
    original_path = os.path.join(TEMP_DIR, f"{file_id}_original.docx")
    logger.info(f"Step 3: Generated unique file ID: '{file_id}'. Target temporary path: '{original_path}'.")
    
    # Save original document to temporary storage
    logger.info("Step 4: Writing uploaded file streams to local temporary disk storage...")
    try:
        with open(original_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        # Verify that the file was written and exists
        if not os.path.exists(original_path) or os.path.getsize(original_path) == 0:
            raise IOError("Temporary file was not written successfully or is empty.")
        logger.info(f"Step 5: File saved successfully. Size: {os.path.getsize(original_path)} bytes.")
    except Exception as e:
        logger.error(f"Failed to write uploaded file to disk: {e}")
        # Clean up just in case
        if os.path.exists(original_path):
            try:
                os.remove(original_path)
            except Exception:
                pass
        raise HTTPException(
            status_code=500, 
            detail="Could not save the uploaded document on server."
        )

    # All steps after this must clean up original_path if they fail
    try:
        # Validate if python-docx can open/parse the file
        logger.info("Step 6: Opening saved file with python-docx for formatting validation...")
        try:
            import docx
            doc = docx.Document(original_path)
            # Count paragraphs to ensure document structure is valid
            para_count = len(doc.paragraphs)
            logger.info(f"Step 7: python-docx validation passed. Document parsed successfully with {para_count} paragraphs.")
        except Exception as e:
            logger.error(f"python-docx parsing validation failed: {e}")
            raise ValueError("The file is corrupted or is not a valid Microsoft Word (.docx) file.")
            
        # Extract text for metrics calculation
        logger.info("Step 8: Extracting document text for initial metrics calculation...")
        try:
            full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            logger.info(f"Step 9: Text extraction completed. Extracted text length: {len(full_text)} characters.")
        except Exception as e:
            logger.error(f"Failed to extract text from document paragraphs: {e}")
            raise ValueError("Failed to read text paragraphs from the document.")

        logger.info("Step 10: Invoking calculate_metrics helper...")
        try:
            before_metrics = calculate_metrics(full_text)
            logger.info("Step 11: Metrics calculation completed successfully.")
        except Exception as e:
            logger.error(f"Failed to calculate metrics: {e}")
            raise ValueError("Failed to analyze document metrics.")
            
        # Construct a meaningful document summary
        logger.info("Step 12: Constructing text-based document summary description...")
        summary = (
            f"This document contains {before_metrics['word_count']} words spread across "
            f"{before_metrics['sentence_count']} sentences. It is written at a {before_metrics['grade_level']} "
            f"level, and has an estimated AI score of {before_metrics['ai_score']}%."
        )
        logger.info("Step 13: Summary generation completed successfully.")
        
        logger.info(f"Step 14: File upload and processing workflow finished successfully for file ID: {file_id}. Word count: {before_metrics['word_count']}.")
        para_risks = calculate_paragraph_risks(full_text)
        return {
            "file_id": file_id,
            "metrics": before_metrics,
            "summary": summary,
            "paragraph_risks": para_risks
        }
    except Exception as e:
        logger.error(f"Error occurred during post-upload file processing: {e}")
        # Clean up temporary file to prevent leaking disk space
        if os.path.exists(original_path):
            try:
                os.remove(original_path)
                logger.info(f"Cleaned up temporary file: {original_path} due to processing error.")
            except Exception as rm_err:
                logger.error(f"Failed to remove temporary file {original_path}: {rm_err}")
                
        # Raise appropriate HTTP response without leaking internal traceback detail
        if isinstance(e, ValueError):
            raise HTTPException(status_code=400, detail=str(e))
        else:
            raise HTTPException(status_code=500, detail="Internal server error during document processing.")

# 3. POST /api/humanize (SSE stream)
@app.post("/api/humanize")
async def humanize_endpoint(request: HumanizeRequest):
    file_id = request.file_id
    api_key = request.api_key
    
    original_path = os.path.join(TEMP_DIR, f"{file_id}_original.docx")
    if not os.path.exists(original_path):
        logger.warning(f"Requested humanization for non-existent file: {file_id}")
        raise HTTPException(status_code=404, detail="Original document file not found.")
        
    async def sse_event_generator():
        try:
            logger.info(f"Starting parallel chunk-by-chunk humanization for document: {file_id}")
            
            # Non-blocking extraction of chunks
            chunks = await asyncio.to_thread(DocxProcessor.extract_chunks, original_path)
            total_chunks = len(chunks)
            logger.info(f"Successfully extracted {total_chunks} chunks from document {file_id}")
            
            # If document is empty, finalize immediately
            if total_chunks == 0:
                humanized_path = os.path.join(TEMP_DIR, f"{file_id}_humanized.docx")
                await asyncio.to_thread(shutil.copyfile, original_path, humanized_path)
                
                empty_metrics = {
                    "word_count": 0,
                    "character_count": 0,
                    "sentence_count": 0,
                    "syllable_count": 0,
                    "readability_score": 0.0,
                    "ai_score": 0.0,
                    "human_score": 100.0,
                    "grade_level": "N/A",
                    "vocabulary_richness": 0.0,
                    "estimated_reading_time_minutes": 0,
                    "estimated_reading_time_seconds": 0
                }
                
                logger.info(f"Finished processing empty document {file_id}")
                yield f"data: {json.dumps({'status': 'completed', 'file_id': file_id, 'metrics': empty_metrics, 'message': 'The document contains no editable paragraphs.'})}\n\n"
                return
                
            # Create a queue to coordinate SSE events from parallel tasks
            event_queue = asyncio.Queue()
            
            # Create a semaphore to cap concurrent Gemini API requests to 3 to stay within rate limits
            api_semaphore = asyncio.Semaphore(3)
            
            # Completed chunks counter with a lock for thread-safe operations
            completed_chunks = 0
            completed_chunks_lock = asyncio.Lock()
            
            async def process_chunk_task(chunk_data, chunk_idx):
                nonlocal completed_chunks
                c_index = chunk_data["index"]
                orig_text = chunk_data["text"]
                
                # Check for empty paragraph/text cleanly to skip external API calls
                if not orig_text.strip():
                    async with completed_chunks_lock:
                        completed_chunks += 1
                        progress_pct = int((completed_chunks / total_chunks) * 100)
                    
                    event = {
                        "status": "processing",
                        "current_chunk": chunk_idx + 1,
                        "total_chunks": total_chunks,
                        "progress_percentage": progress_pct,
                        "original_text": orig_text,
                        "humanized_text": orig_text,
                        "chunk_index": c_index
                    }
                    await event_queue.put(event)
                    return c_index, orig_text
                
                # Run the actual humanization using the concurrency limiting semaphore
                async with api_semaphore:
                    # Notify about starting API processing
                    await event_queue.put({
                        "status": "stage",
                        "current_chunk": chunk_idx + 1,
                        "total_chunks": total_chunks,
                        "message": f"Humanizing chunk {chunk_idx+1}/{total_chunks} via Gemini API..."
                    })
                    
                    try:
                        # Call the blocking humanizer function in a separate thread to keep the event loop responsive
                        start_time = asyncio.get_event_loop().time()
                        hum_text = await humanize_text(orig_text, api_key, preset=request.preset, ngram_shield=request.ngram_shield)
                        duration = asyncio.get_event_loop().time() - start_time
                        logger.info(f"Successfully humanized chunk {chunk_idx+1}/{total_chunks} (index {c_index}) in {duration:.2f}s")
                        
                        # Notify about entering post-processing
                        await event_queue.put({
                            "status": "stage",
                            "current_chunk": chunk_idx + 1,
                            "total_chunks": total_chunks,
                            "message": "Applying anti-detection post-processing..."
                        })
                        
                        async with completed_chunks_lock:
                            completed_chunks += 1
                            progress_pct = int((completed_chunks / total_chunks) * 100)
                            
                        # Put the successful humanization result on the queue
                        event = {
                            "status": "processing",
                            "current_chunk": chunk_idx + 1,
                            "total_chunks": total_chunks,
                            "progress_percentage": progress_pct,
                            "original_text": orig_text,
                            "humanized_text": hum_text,
                            "chunk_index": c_index
                        }
                        await event_queue.put(event)
                        return c_index, hum_text
                        
                    except RateLimitError as e:
                        logger.error(f"Rate limit exceeded during chunk {chunk_idx+1} humanization: {e}")
                        await event_queue.put({
                            "status": "error",
                            "message": "Gemini API rate limit exceeded. Please check your quota/API key rate limit, or try again later."
                        })
                        raise
                    except GeminiAPIError as e:
                        logger.error(f"Gemini API error during chunk {chunk_idx+1} humanization: {e}")
                        await event_queue.put({
                            "status": "error",
                            "message": f"Gemini API Error: {str(e)}"
                        })
                        raise
                    except HumanizerError as e:
                        logger.error(f"Humanization connection or timeout error during chunk {chunk_idx+1}: {e}")
                        await event_queue.put({
                            "status": "error",
                            "message": f"Connection/Humanization Error: {str(e)}"
                        })
                        raise
                    except Exception as e:
                        logger.exception(f"Unexpected error during chunk {chunk_idx+1} humanization: {e}")
                        await event_queue.put({
                            "status": "error",
                            "message": f"An unexpected error occurred: {str(e)}"
                        })
                        raise
            
            # Spawn processing tasks for all chunks concurrently
            tasks = [
                asyncio.create_task(process_chunk_task(c, i))
                for i, c in enumerate(chunks)
            ]
            
            # Run a monitor task to wait for completion and put a None sentinel in the event queue
            async def monitor_tasks():
                try:
                    await asyncio.gather(*tasks, return_exceptions=True)
                except Exception as e:
                    logger.error(f"Monitor caught error: {e}")
                finally:
                    await event_queue.put(None)
                    
            monitor_job = asyncio.create_task(monitor_tasks())
            
            # Read from the event queue and stream to the SSE client in real time
            updated_chunks = {}
            has_error = False
            
            while True:
                event = await event_queue.get()
                if event is None:
                    break
                    
                if event.get("status") == "error":
                    has_error = True
                    yield f"data: {json.dumps(event)}\n\n"
                    # Cancel any outstanding chunk tasks
                    for t in tasks:
                        if not t.done():
                            t.cancel()
                    break
                
                if event.get("status") == "processing":
                    chunk_index = event["chunk_index"]
                    humanized_text = event["humanized_text"]
                    updated_chunks[chunk_index] = humanized_text
                    
                    # Yield standard frontend event
                    frontend_event = {
                        "status": "processing",
                        "current_chunk": event["current_chunk"],
                        "total_chunks": event["total_chunks"],
                        "progress_percentage": event["progress_percentage"],
                        "original_text": event["original_text"],
                        "humanized_text": event["humanized_text"]
                    }
                    yield f"data: {json.dumps(frontend_event)}\n\n"
                else:
                    yield f"data: {json.dumps(event)}\n\n"
            
            # Wait for monitor job to clean up fully
            await monitor_job
            
            if has_error:
                return
                
            # Reconstruct humanized docx in a separate thread
            humanized_path = os.path.join(TEMP_DIR, f"{file_id}_humanized.docx")
            logger.info(f"Rebuilding docx and saving to {humanized_path}")
            await asyncio.to_thread(DocxProcessor.replace_chunks, original_path, humanized_path, updated_chunks)
            
            # Calculate final metrics from humanized file in a separate thread
            import docx
            o_doc = await asyncio.to_thread(docx.Document, original_path)
            h_doc = await asyncio.to_thread(docx.Document, humanized_path)
            
            full_original_text = "\n\n".join([p.text for p in o_doc.paragraphs if p.text.strip()])
            full_humanized_text = "\n\n".join([p.text for p in h_doc.paragraphs if p.text.strip()])
            
            # CPU-bound metrics calculation run in a separate thread
            after_metrics = await asyncio.to_thread(calculate_metrics, full_humanized_text)
            para_risks = await asyncio.to_thread(calculate_paragraph_risks, full_humanized_text)
            
            # Yield completed event
            final_data = {
                "status": "completed",
                "file_id": file_id,
                "metrics": after_metrics,
                "original_text": full_original_text,
                "humanized_text": full_humanized_text,
                "paragraph_risks": para_risks,
                "message": "Humanization completed successfully."
            }
            logger.info(f"Successfully completed humanizing file: {file_id}")
            yield f"data: {json.dumps(final_data)}\n\n"
            
        except Exception as e:
            logger.exception(f"Unexpected error while humanizing file {file_id}: {e}")
            error_data = {
                "status": "error",
                "message": f"Processing error: {str(e)}"
            }
            yield f"data: {json.dumps(error_data)}\n\n"
            
    return StreamingResponse(sse_event_generator(), media_type="text/event-stream")

# Schema for document update request
class UpdateDocumentRequest(BaseModel):
    humanized_text: str

@app.post("/api/update/{file_id}")
async def update_document(file_id: str, request: UpdateDocumentRequest):
    logger.info(f"Update document request received for: {file_id}")
    original_path = os.path.join(TEMP_DIR, f"{file_id}_original.docx")
    humanized_path = os.path.join(TEMP_DIR, f"{file_id}_humanized.docx")
    
    if not os.path.exists(original_path):
        logger.warning(f"Update rejected. Original document file not found: {file_id}")
        raise HTTPException(status_code=404, detail="Original document file not found.")
        
    try:
        # Import DocxProcessor inside the endpoint if not globally imported
        try:
            from backend.docx_processor import DocxProcessor
        except ImportError:
            from docx_processor import DocxProcessor
            
        processor = DocxProcessor()
        await asyncio.to_thread(processor.load, original_path)
        
        # Split user's text into paragraphs
        new_paragraphs = [p.strip() for p in request.humanized_text.split("\n\n")]
        new_paragraphs = [p for p in new_paragraphs if p]
        
        # Find original non-empty paragraph indices
        non_empty_indices = []
        for idx, p in enumerate(processor.paragraphs):
            if p.text.strip():
                non_empty_indices.append(idx)
                
        # Perform replacements preserving formatting in-place
        for idx, orig_idx in enumerate(non_empty_indices):
            if idx < len(new_paragraphs):
                await asyncio.to_thread(processor.replace_paragraph_text, orig_idx, new_paragraphs[idx])
                
        # Save updated document
        await asyncio.to_thread(processor.save, humanized_path)
        
        # Recalculate metrics for updated document
        full_text = "\n\n".join(new_paragraphs)
        updated_metrics = await asyncio.to_thread(calculate_metrics, full_text)
        para_risks = await asyncio.to_thread(calculate_paragraph_risks, full_text)
        
        logger.info(f"Successfully updated and saved document: {file_id}")
        return {
            "status": "success",
            "metrics": updated_metrics,
            "paragraph_risks": para_risks
        }
    except Exception as e:
        logger.exception(f"Failed to update document {file_id}: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to apply user edits to document: {str(e)}")

# 4. GET /api/download/{file_id}
@app.get("/api/download/{file_id}")
async def download_file(file_id: str):
    logger.info(f"Download request received for: {file_id}")
    humanized_path = os.path.join(TEMP_DIR, f"{file_id}_humanized.docx")
    
    if not os.path.exists(humanized_path):
        original_path = os.path.join(TEMP_DIR, f"{file_id}_original.docx")
        if os.path.exists(original_path):
            logger.warning(f"Download rejected. Document {file_id} has not been humanized yet.")
            raise HTTPException(
                status_code=400, 
                detail="This document exists but has not been humanized yet. Please humanize it first."
            )
        logger.warning(f"Download rejected. File ID not found: {file_id}")
        raise HTTPException(
            status_code=404, 
            detail="File ID not found or expired."
        )
        
    return FileResponse(
        path=humanized_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="humanized_document.docx"
    )

# ============================================================
# 5. STATIC FILE SERVING (Cloud Run Production Mode)
# ============================================================
# In production (Cloud Run), the built React frontend is copied
# into /app/static/ by the Dockerfile. This mounts it so FastAPI
# serves both the API and the frontend from a single container.
STATIC_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "static")
if os.path.isdir(STATIC_DIR):
    logger.info(f"Production mode: Serving static frontend from {STATIC_DIR}")
    
    # SPA catch-all: serve index.html for any non-API, non-file route
    @app.get("/{full_path:path}")
    async def serve_spa(request: Request, full_path: str):
        # If it's an API route, skip (shouldn't reach here due to ordering)
        if full_path.startswith("api/"):
            raise HTTPException(status_code=404, detail="API endpoint not found.")
        
        # Try to serve the exact static file first
        file_path = os.path.join(STATIC_DIR, full_path)
        if full_path and os.path.isfile(file_path):
            return FileResponse(file_path)
        
        # Fallback to index.html for SPA client-side routing
        index_path = os.path.join(STATIC_DIR, "index.html")
        if os.path.isfile(index_path):
            return FileResponse(index_path, media_type="text/html")
        
        raise HTTPException(status_code=404, detail="Resource not found.")
else:
    logger.info("Development mode: No static directory found. Frontend served separately.")
