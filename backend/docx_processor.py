import os
import logging
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table

logger = logging.getLogger("backend.docx_processor")

class DocxProcessor:
    """
    A utility class for processing and rephrasing Microsoft Word (.docx) documents.
    It preserves the visual and structural formatting (styles, bullet points, list items,
    images, and drawings) by safely replacing run-level text in-place.
    """

    def __init__(self):
        self.doc = None
        self.paragraphs = []

    def load(self, file_path: str):
        """
        Loads a .docx file using python-docx and extracts all paragraph elements.
        
        Args:
            file_path: The absolute or relative path to the .docx file.
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        self.doc = Document(file_path)
        self.paragraphs = []
        self._extract_all_paragraphs()

    def _extract_all_paragraphs(self):
        """
        Recursively extracts all paragraph objects (including body text, headings, list items,
        and table cell paragraphs) in their exact visual/structural order from the document body.
        """
        self.paragraphs = []
        if self.doc is None:
            return

        def traverse_element(element):
            tag = element.tag.split('}')[-1]
            if tag == 'p':
                self.paragraphs.append(Paragraph(element, self.doc))
            elif tag == 'tbl':
                table = Table(element, self.doc)
                for row in table.rows:
                    for cell in row.cells:
                        for child in cell._element:
                            traverse_element(child)

        for child in self.doc.element.body:
            traverse_element(child)

    def get_paragraphs(self):
        """
        Returns the list of all extracted paragraph objects.
        
        Returns:
            A list of docx.text.paragraph.Paragraph objects.
        """
        return self.paragraphs

    def group_into_chunks(self, target_words: int = 700, min_words: int = 600, max_words: int = 800):
        """
        Groups extracted paragraph/cell text into logical chunks of ~600-800 words,
        ignoring empty paragraphs/cells, and returning these chunks along with a
        mapping back to their original paragraph index.
        
        Args:
            target_words: The ideal number of words per chunk.
            min_words: The minimum number of words allowed before closing a chunk.
            max_words: The maximum number of words allowed in a chunk.
            
        Returns:
            List of Dicts: Each dictionary contains:
                - 'chunk_index': The sequential ID of the chunk (0-indexed).
                - 'text': The concatenated text of the paragraphs/cells in the chunk,
                          separated by newlines.
                - 'mapping': A list of the original paragraph indices corresponding to the
                             paragraphs/cells that make up this chunk.
        """
        logger.info(f"group_into_chunks: Starting grouping. target_words={target_words}, min_words={min_words}, max_words={max_words}")
        chunks = []
        current_chunk_paragraphs = []
        current_word_count = 0
        chunk_idx = 0

        for idx, p in enumerate(self.paragraphs):
            text = p.text.strip()
            if not text:
                continue
            
            words = text.split()
            num_words = len(words)
            
            # If adding this paragraph would exceed max_words, close the current chunk first.
            if current_word_count + num_words > max_words and current_chunk_paragraphs:
                mapped_indices = [original_idx for original_idx, _ in current_chunk_paragraphs]
                logger.info(
                    f"group_into_chunks: Closing Chunk {chunk_idx} (word count {current_word_count}) "
                    f"prior to paragraph {idx} of {num_words} words (would exceed max_words={max_words}). "
                    f"Paragraphs in chunk: {mapped_indices}"
                )
                chunks.append({
                    "chunk_index": chunk_idx,
                    "text": "\n".join(t for _, t in current_chunk_paragraphs),
                    "mapping": mapped_indices
                })
                chunk_idx += 1
                current_chunk_paragraphs = []
                current_word_count = 0
            
            current_chunk_paragraphs.append((idx, text))
            current_word_count += num_words
            
            # If we've reached the target word count, close the chunk.
            if current_word_count >= target_words:
                mapped_indices = [original_idx for original_idx, _ in current_chunk_paragraphs]
                logger.info(
                    f"group_into_chunks: Closing Chunk {chunk_idx} (word count {current_word_count}) "
                    f"as it met/exceeded target_words={target_words}. Paragraphs in chunk: {mapped_indices}"
                )
                chunks.append({
                    "chunk_index": chunk_idx,
                    "text": "\n".join(t for _, t in current_chunk_paragraphs),
                    "mapping": mapped_indices
                })
                chunk_idx += 1
                current_chunk_paragraphs = []
                current_word_count = 0

        # Close the last remaining chunk
        if current_chunk_paragraphs:
            mapped_indices = [original_idx for original_idx, _ in current_chunk_paragraphs]
            logger.info(
                f"group_into_chunks: Closing final Chunk {chunk_idx} (word count {current_word_count}). "
                f"Paragraphs in chunk: {mapped_indices}"
            )
            chunks.append({
                "chunk_index": chunk_idx,
                "text": "\n".join(t for _, t in current_chunk_paragraphs),
                "mapping": mapped_indices
            })
            
        logger.info(f"group_into_chunks: Done. Generated {len(chunks)} chunks.")
        return chunks

    def replace_paragraph_text(self, index: int, new_text: str):
        """
        Replaces the text of the paragraph at the specified index in-place.
        To perfectly preserve formatting, styles, bullet points, numbers, drawings, and images,
        it clears text nodes from all runs in the paragraph except the first run, and sets
        the first run's text to the new rephrased content.
        
        Args:
            index: The index of the paragraph to replace.
            new_text: The new rephrased string to assign.
        """
        if index < 0 or index >= len(self.paragraphs):
            raise IndexError(f"Paragraph index {index} is out of range.")
            
        paragraph = self.paragraphs[index]
        runs = paragraph.runs
        
        if not runs:
            # If there are no runs, add one
            paragraph.add_run(new_text)
        else:
            # Set the first run's text safely
            self._set_run_text_safe(runs[0], new_text)
            # Clear text of all subsequent runs safely to preserve run properties/images
            for run in runs[1:]:
                self._clear_run_text(run)

    def _set_run_text_safe(self, run, text: str):
        """
        Safely sets the text of a run by updating the first w:t element
        and removing any other w:t elements, preserving properties/drawings.
        """
        t_elements = run._element.xpath('w:t')
        if not t_elements:
            run.text = text
        else:
            t_elements[0].text = text
            for t in t_elements[1:]:
                run._element.remove(t)

    def _clear_run_text(self, run):
        """
        Safely clears all text from a run by removing all w:t elements.
        """
        t_elements = run._element.xpath('w:t')
        for t in t_elements:
            run._element.remove(t)

    def save(self, new_path: str):
        """
        Saves the modified document to a new path.
        
        Args:
            new_path: The destination file path.
        """
        if self.doc is None:
            raise ValueError("No document loaded to save.")
        
        # Ensure parent directories exist
        parent_dir = os.path.dirname(new_path)
        if parent_dir and not os.path.exists(parent_dir):
            os.makedirs(parent_dir, exist_ok=True)
            
        self.doc.save(new_path)

    @classmethod
    def extract_chunks(cls, file_path: str, target_words: int = 700, min_words: int = 600, max_words: int = 800):
        """
        Class method to load a docx file and extract logical chunks of text.
        
        Args:
            file_path: The path to the .docx file.
            target_words: The target word count for chunks.
            min_words: The minimum word count.
            max_words: The maximum word count.
            
        Returns:
            List of chunks with 'index', 'text', and 'mapping'.
        """
        processor = cls()
        processor.load(file_path)
        chunks = processor.group_into_chunks(target_words=target_words, min_words=min_words, max_words=max_words)
        # Add the 'index' key as an alias to 'chunk_index' to maintain compatibility with backend/main.py
        for chunk in chunks:
            chunk["index"] = chunk["chunk_index"]
        return chunks

    @classmethod
    def replace_chunks(cls, original_path: str, new_path: str, updated_chunks: dict):
        """
        Class method to load the original document, replace the specified chunks,
        and save the final document to new_path.
        
        Args:
            original_path: The path to the original .docx file.
            new_path: The path to save the modified docx.
            updated_chunks: A dictionary mapping chunk index to its new humanized text.
        """
        processor = cls()
        processor.load(original_path)
        
        # Get chunks to resolve mapping of chunk index to paragraph indices
        chunks = processor.group_into_chunks()
        chunks_by_index = {c["chunk_index"]: c for c in chunks}
        
        logger.info(f"replace_chunks: Beginning chunk replacement. Total chunks to apply: {len(updated_chunks)}")
        
        for chunk_idx_str, humanized_text in updated_chunks.items():
            # Convert key to int in case it is serialized as a string in JSON or SSE parsing
            try:
                chunk_idx = int(chunk_idx_str)
            except ValueError:
                logger.error(f"replace_chunks: Invalid chunk index type: {chunk_idx_str}")
                continue
                
            if chunk_idx not in chunks_by_index:
                logger.warning(f"replace_chunks: Chunk index {chunk_idx} not found in document. Skipping replacement.")
                continue
                
            chunk = chunks_by_index[chunk_idx]
            mapping = chunk["mapping"]
            num_expected = len(mapping)
            
            logger.info(f"replace_chunks: Processing Chunk {chunk_idx} (expects {num_expected} paragraphs, maps to: {mapping})")
            
            # Split the humanized text into paragraphs. We support splitting by \n\n or \n.
            paras_double = [p.strip() for p in humanized_text.split("\n\n")]
            paras_double = [p for p in paras_double if p]
            
            paras_single = [p.strip() for p in humanized_text.split("\n")]
            paras_single = [p for p in paras_single if p]
            
            if len(paras_double) == num_expected:
                replacement_paras = paras_double
                logger.info(f"replace_chunks: Chunk {chunk_idx} aligned perfectly using double-newlines.")
            elif len(paras_single) == num_expected:
                replacement_paras = paras_single
                logger.info(f"replace_chunks: Chunk {chunk_idx} aligned perfectly using single-newlines.")
            else:
                # Sequential fallback alignment logic when paragraph count differs from mapping size
                logger.warning(
                    f"replace_chunks: Paragraph count mismatch for Chunk {chunk_idx}. "
                    f"Expected: {num_expected}. Double-split got {len(paras_double)}, Single-split got {len(paras_single)}. "
                    f"Applying sequential fallback alignment."
                )
                
                # Default to single-split as it keeps lines distinct
                replacement_paras = paras_single
                
                if len(replacement_paras) < num_expected:
                    logger.info(f"replace_chunks: Chunk {chunk_idx} got fewer paragraphs ({len(replacement_paras)}). Padding with original text.")
                    for i in range(len(replacement_paras), num_expected):
                        orig_idx = mapping[i]
                        replacement_paras.append(processor.paragraphs[orig_idx].text)
                elif len(replacement_paras) > num_expected:
                    logger.info(f"replace_chunks: Chunk {chunk_idx} got more paragraphs ({len(replacement_paras)}). Merging extras.")
                    keep = replacement_paras[:num_expected - 1]
                    extra = replacement_paras[num_expected - 1:]
                    keep.append("\n\n".join(extra))
                    replacement_paras = keep
            
            # Perform replacement in-place
            for orig_idx, new_para_text in zip(mapping, replacement_paras):
                logger.info(f"replace_chunks: Replacing Paragraph {orig_idx} with humanized text of length {len(new_para_text)}")
                processor.replace_paragraph_text(orig_idx, new_para_text)
                
        logger.info(f"replace_chunks: Saving modified document to {new_path}")
        processor.save(new_path)
