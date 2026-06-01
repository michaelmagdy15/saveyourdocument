import os
import unittest
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from backend.docx_processor import DocxProcessor

class TestDocxProcessor(unittest.TestCase):
    def setUp(self):
        self.test_dir = os.path.dirname(os.path.abspath(__file__))
        self.input_path = os.path.join(self.test_dir, "sample_test.docx")
        self.output_path = os.path.join(self.test_dir, "output_test.docx")
        
        # Ensure test directory exists
        os.makedirs(self.test_dir, exist_ok=True)
        
        # Create a sample docx file with complex structures
        doc = Document()
        
        # 1. Add Heading
        doc.add_heading("Main Title", level=1)
        
        # 2. Add standard paragraphs with rich formatting
        p1 = doc.add_paragraph()
        r1 = p1.add_run("This is the first paragraph. ")
        r2 = p1.add_run("It has bold text.")
        r2.bold = True
        r3 = p1.add_run(" And some italic text.")
        r3.italic = True
        
        # 3. Add an empty paragraph
        doc.add_paragraph("")
        
        # 4. Add list items
        doc.add_paragraph("First item of numbered list", style='List Number')
        doc.add_paragraph("Second item of numbered list", style='List Number')
        doc.add_paragraph("First item of bullet list", style='List Bullet')
        doc.add_paragraph("Second item of bullet list", style='List Bullet')
        
        # 5. Add a table
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).paragraphs[0].text = "Cell 1,1 text"
        table.cell(0, 1).paragraphs[0].text = "Cell 1,2 text"
        
        # Empty cell paragraph
        table.cell(1, 0).paragraphs[0].text = ""
        
        # Multi-paragraph cell
        cell_1_1 = table.cell(1, 1)
        cell_1_1.paragraphs[0].text = "Cell 2,2 first paragraph"
        p_cell_2 = cell_1_1.add_paragraph()
        p_cell_2.add_run("Cell 2,2 second paragraph with ")
        r_bold = p_cell_2.add_run("bold part")
        r_bold.bold = True
        
        # 6. Add large paragraphs to test chunking (needs ~1200 words to make multiple chunks)
        # We will add 3 paragraphs of ~250 words each
        word_block = "word " * 250
        doc.add_paragraph(f"Block 1: {word_block}")
        doc.add_paragraph(f"Block 2: {word_block}")
        doc.add_paragraph(f"Block 3: {word_block}")
        
        doc.save(self.input_path)

    def tearDown(self):
        # Clean up files
        for path in [self.input_path, self.output_path]:
            if os.path.exists(path):
                try:
                    os.remove(path)
                except OSError:
                    pass

    def test_docx_processor_flow(self):
        processor = DocxProcessor()
        
        # Test 1: Load file
        processor.load(self.input_path)
        paragraphs = processor.get_paragraphs()
        
        # Let's verify number of extracted paragraphs
        # Heading (1) + p1 (1) + empty p (1) + list items (4) + Table cells (4 cells with paragraphs: cell 0,0: 1 p; cell 0,1: 1 p; cell 1,0: 1 p; cell 1,1: 2 p; total table paragraphs = 5) + 3 large blocks (3)
        # Total expected paragraphs = 1 + 1 + 1 + 4 + 5 + 3 = 15
        self.assertEqual(len(paragraphs), 15)
        
        # Verify specific texts
        self.assertEqual(paragraphs[0].text, "Main Title")
        self.assertTrue("italic text" in paragraphs[1].text)
        self.assertEqual(paragraphs[2].text, "")
        self.assertEqual(paragraphs[3].text, "First item of numbered list")
        self.assertEqual(paragraphs[7].text, "Cell 1,1 text")
        self.assertEqual(paragraphs[9].text, "") # cell 1,0 is empty
        self.assertEqual(paragraphs[11].text, "Cell 2,2 second paragraph with bold part")
        
        # Test 2: Chunking logic
        # Our blocks 1, 2, 3 have ~250 words each (total ~750 words + other text).
        # We can configure grouping parameters to test chunk boundary splitting.
        chunks = processor.group_into_chunks(target_words=300, min_words=200, max_words=400)
        
        print("\n--- Chunking Results ---")
        for chunk in chunks:
            words = len(chunk['text'].split())
            print(f"Chunk {chunk['chunk_index']}: words={words}, mapping={chunk['mapping']}")
        
        # We expect multiple chunks because target_words=300 and total words is >750.
        self.assertGreater(len(chunks), 1)
        
        # Verify that empty paragraphs are ignored from chunk mappings
        for chunk in chunks:
            # Paragraph at index 2 is empty, index 9 is empty. They should not be in any mapping.
            self.assertNotIn(2, chunk['mapping'])
            self.assertNotIn(9, chunk['mapping'])
            
        # Test 3: Replace paragraph text in-place while keeping properties
        # Rephrase a standard paragraph with bold/italic runs
        processor.replace_paragraph_text(1, "This is a rephrased paragraph with bold and italic details.")
        
        # Rephrase a list item and table cells
        processor.replace_paragraph_text(3, "Rephrased numbered list item 1")
        processor.replace_paragraph_text(7, "Rephrased cell text")
        
        # Test 4: Save and verify modified document
        processor.save(self.output_path)
        
        # Load the saved file to verify properties
        new_doc = Document(self.output_path)
        new_paragraphs = []
        
        # Extract body paragraphs using recursive traversal to match processor's structure
        def traverse_element(element):
            tag = element.tag.split('}')[-1]
            if tag == 'p':
                new_paragraphs.append(Paragraph(element, new_doc))
            elif tag == 'tbl':
                table = Table(element, new_doc)
                for row in table.rows:
                    for cell in row.cells:
                        for child in cell._element:
                            traverse_element(child)

        for child in new_doc.element.body:
            traverse_element(child)
            
        self.assertEqual(len(new_paragraphs), 15)
        
        # Verify text replacements
        self.assertEqual(new_paragraphs[1].text, "This is a rephrased paragraph with bold and italic details.")
        self.assertEqual(new_paragraphs[3].text, "Rephrased numbered list item 1")
        self.assertEqual(new_paragraphs[7].text, "Rephrased cell text")
        
        # Verify formatting preservation
        # For new_paragraphs[1] (originally p1), it had three runs: "This is the first paragraph. ", "It has bold text.", " And some italic text."
        # After replace_paragraph_text, the new text is distributed proportionally across non-empty runs
        # Let's inspect:
        p1_runs = new_paragraphs[1].runs
        self.assertEqual(p1_runs[0].text.strip(), "This is a rephrased")
        self.assertEqual(p1_runs[1].text.strip(), "paragraph with bold")
        self.assertEqual(p1_runs[2].text.strip(), "and italic details.")
            
        # Check list item style remains
        self.assertEqual(new_paragraphs[3].style.name, 'List Number')
        self.assertEqual(new_paragraphs[5].style.name, 'List Bullet')
        
        print("\nAll tests passed successfully!")

if __name__ == '__main__':
    unittest.main()
