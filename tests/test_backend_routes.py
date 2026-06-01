import os
import shutil
import unittest
import json
from unittest.mock import patch, MagicMock
from fastapi.testclient import TestClient
from docx import Document

# Import backend application
try:
    from backend.main import app, TEMP_DIR
except ImportError:
    from main import app, TEMP_DIR

class TestBackendRoutes(unittest.TestCase):
    def setUp(self):
        self.client = TestClient(app)
        self.test_dir = os.path.dirname(os.path.abspath(__file__))
        
        # Paths for temporary test documents
        self.valid_docx_path = os.path.join(self.test_dir, "route_test_valid.docx")
        self.invalid_docx_path = os.path.join(self.test_dir, "route_test_invalid.docx")
        self.invalid_ext_path = os.path.join(self.test_dir, "route_test.txt")
        
        # 1. Create a valid docx file
        doc = Document()
        doc.add_heading("French Language Test Document", level=1)
        doc.add_paragraph("Tout d'abord, nous devons examiner le problème en détail.")
        doc.add_paragraph("De plus, il convient de noter que l'intelligence artificielle progresse rapidement.")
        doc.save(self.valid_docx_path)
        
        # 2. Create an invalid/corrupted docx file (just plain text with docx extension)
        with open(self.invalid_docx_path, "w", encoding="utf-8") as f:
            f.write("This is not a real zip-based docx file, it's just plain text.")
            
        # 3. Create a file with a disallowed extension
        with open(self.invalid_ext_path, "w", encoding="utf-8") as f:
            f.write("Some text file contents.")

    def tearDown(self):
        # Clean up files in tests/ folder
        for path in [self.valid_docx_path, self.invalid_docx_path, self.invalid_ext_path]:
            if os.path.exists(path):
                try:
                    os.remove(path)
                except OSError:
                    pass
                    
        # Also clean up any lingering test files in TEMP_DIR
        if os.path.exists(TEMP_DIR):
            for filename in os.listdir(TEMP_DIR):
                if filename.endswith(".docx"):
                    try:
                        os.remove(os.path.join(TEMP_DIR, filename))
                    except OSError:
                        pass

    def test_root_endpoint(self):
        response = self.client.get("/")
        self.assertEqual(response.status_code, 200)
        self.assertIn("status", response.json())
        self.assertEqual(response.json()["status"], "healthy")

    def test_upload_invalid_extension(self):
        # Verify extension rejection
        with open(self.invalid_ext_path, "rb") as f:
            response = self.client.post(
                "/api/upload",
                files={"file": ("route_test.txt", f, "text/plain")}
            )
        self.assertEqual(response.status_code, 400)
        self.assertEqual(response.json()["detail"], "Only word documents (.docx) are supported.")

    def test_upload_corrupted_docx(self):
        # Verify file corruption validation and file cleanup on error
        initial_temp_files = os.listdir(TEMP_DIR)
        
        with open(self.invalid_docx_path, "rb") as f:
            response = self.client.post(
                "/api/upload",
                files={"file": ("route_test_invalid.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
            )
        self.assertEqual(response.status_code, 400)
        self.assertEqual(response.json()["detail"], "The file is corrupted or is not a valid Microsoft Word (.docx) file.")
        
        # Verify that the corrupted file was cleaned up and not left inside temp_uploads
        final_temp_files = os.listdir(TEMP_DIR)
        self.assertEqual(len(initial_temp_files), len(final_temp_files))

    def test_upload_valid_docx_metrics(self):
        # Verify successful upload and flat metrics formatting
        with open(self.valid_docx_path, "rb") as f:
            response = self.client.post(
                "/api/upload",
                files={"file": ("route_test_valid.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
            )
        self.assertEqual(response.status_code, 200)
        
        data = response.json()
        self.assertIn("file_id", data)
        self.assertIn("summary", data)
        self.assertIn("metrics", data)
        
        # Verify flat metrics schema coordination
        metrics = data["metrics"]
        expected_keys = {
            "word_count", "character_count", "sentence_count", "syllable_count",
            "readability_score", "ai_score", "human_score", "grade_level",
            "vocabulary_richness", "estimated_reading_time_minutes", "estimated_reading_time_seconds"
        }
        for key in expected_keys:
            self.assertIn(key, metrics, f"Metric key '{key}' is missing from returned metrics dictionary.")
            
        # Verify specific calculations
        self.assertGreater(metrics["word_count"], 0)
        self.assertGreater(metrics["character_count"], 0)
        self.assertGreater(metrics["sentence_count"], 0)
        self.assertGreater(metrics["syllable_count"], 0)
        self.assertEqual(metrics["ai_score"] + metrics["human_score"], 100.0)
        self.assertNotEqual(metrics["grade_level"], "N/A")
        
        # Verify that original file persists in TEMP_DIR for subsequent humanization endpoint calls
        file_id = data["file_id"]
        original_temp_file = os.path.join(TEMP_DIR, f"{file_id}_original.docx")
        self.assertTrue(os.path.exists(original_temp_file))

    @patch("backend.main.humanize_text")
    def test_humanize_route_success(self, mock_humanize_text):
        # Setup mock behavior
        mock_humanize_text.return_value = "C'est une version réécrite très humaine et élégante."
        
        # Upload valid file first to get file_id
        with open(self.valid_docx_path, "rb") as f:
            upload_resp = self.client.post(
                "/api/upload",
                files={"file": ("route_test_valid.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
            )
        self.assertEqual(upload_resp.status_code, 200)
        file_id = upload_resp.json()["file_id"]
        
        # Request humanization via SSE route
        response = self.client.post(
            "/api/humanize",
            json={"file_id": file_id, "api_key": "dummy-api-key"}
        )
        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.headers["content-type"], "text/event-stream; charset=utf-8")
        
        # Parse SSE data events
        events = []
        for line in response.iter_lines():
            if line.startswith("data:"):
                event_content = line[len("data:"):].strip()
                events.append(json.loads(event_content))
                
        # We expect at least one processing event and a completed event
        self.assertTrue(len(events) >= 2)
        
        # The last event must be 'completed'
        self.assertEqual(events[-1]["status"], "completed")
        self.assertEqual(events[-1]["file_id"], file_id)
        self.assertIn("metrics", events[-1])
        
        # Check processing events
        processing_events = [e for e in events if e["status"] == "processing"]
        self.assertTrue(len(processing_events) > 0)
        self.assertEqual(processing_events[0]["humanized_text"], "C'est une version réécrite très humaine et élégante.")

    @patch("backend.main.humanize_text")
    def test_humanize_route_rate_limit(self, mock_humanize_text):
        # Setup mock behavior to raise RateLimitError
        from backend.humanizer import RateLimitError
        mock_humanize_text.side_effect = RateLimitError(429, "Rate limit exceeded.")
        
        # Upload valid file first to get file_id
        with open(self.valid_docx_path, "rb") as f:
            upload_resp = self.client.post(
                "/api/upload",
                files={"file": ("route_test_valid.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
            )
        self.assertEqual(upload_resp.status_code, 200)
        file_id = upload_resp.json()["file_id"]
        
        # Request humanization
        response = self.client.post(
            "/api/humanize",
            json={"file_id": file_id, "api_key": "dummy-api-key"}
        )
        self.assertEqual(response.status_code, 200)
        
        events = []
        for line in response.iter_lines():
            if line.startswith("data:"):
                event_content = line[len("data:"):].strip()
                events.append(json.loads(event_content))
                
        # We expect an error event representing rate limit in the stream
        error_events = [e for e in events if e["status"] == "error"]
        self.assertTrue(len(error_events) > 0)
        self.assertIn("rate limit exceeded", error_events[0]["message"].lower())

    def test_humanize_route_empty_document(self):
        # Create an empty docx document (without paragraphs)
        empty_docx_path = os.path.join(self.test_dir, "route_test_empty.docx")
        doc = Document()
        doc.save(empty_docx_path)
        
        # Upload empty document
        with open(empty_docx_path, "rb") as f:
            upload_resp = self.client.post(
                "/api/upload",
                files={"file": ("route_test_empty.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
            )
        self.assertEqual(upload_resp.status_code, 200)
        file_id = upload_resp.json()["file_id"]
        
        # Clean up the local empty test file
        if os.path.exists(empty_docx_path):
            os.remove(empty_docx_path)
            
        # Request humanization
        response = self.client.post(
            "/api/humanize",
            json={"file_id": file_id, "api_key": "dummy-api-key"}
        )
        self.assertEqual(response.status_code, 200)
        
        events = []
        for line in response.iter_lines():
            if line.startswith("data:"):
                event_content = line[len("data:"):].strip()
                events.append(json.loads(event_content))
                
        self.assertEqual(events[0]["status"], "completed")
        self.assertIn("no editable paragraphs", events[0]["message"].lower())
        self.assertEqual(events[0]["metrics"]["word_count"], 0)

if __name__ == '__main__':
    unittest.main()
